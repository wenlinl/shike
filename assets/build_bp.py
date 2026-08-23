# -*- coding: utf-8 -*-
"""Build the ShiKe (Shike) English business plan DOCX (narrative_proposal preset)."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ACCENT = RGBColor(0x2E, 0x74, 0xB5)      # heading blue
DARK = RGBColor(0x1F, 0x4D, 0x78)        # h3 blue
INK = RGBColor(0x1F, 0x3A, 0x5F)         # cover dark blue
GRAY = RGBColor(0x59, 0x5F, 0x66)
HEADER_FILL = "F4F6F9"


def set_run_font(run, name="Calibri", size=11, color=None, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)


def style_setup(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    rpr = normal.element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    rf.set(qn("w:ascii"), "Calibri")
    rf.set(qn("w:hAnsi"), "Calibri")
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.333

    def tune_heading(name, size, color, before, after):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        rpr = st.element.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts")
            rpr.append(rf)
        rf.set(qn("w:ascii"), "Calibri")
        rf.set(qn("w:hAnsi"), "Calibri")
        pf = st.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.keep_with_next = True
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.1

    tune_heading("Heading 1", 16, ACCENT, 18, 10)
    tune_heading("Heading 2", 13, ACCENT, 12, 6)
    tune_heading("Heading 3", 12, DARK, 8, 4)

    for lst in ("List Bullet", "List Number"):
        st = doc.styles[lst]
        st.font.name = "Calibri"
        st.font.size = Pt(11)
        pf = st.paragraph_format
        pf.left_indent = Inches(0.375)
        pf.first_line_indent = Inches(-0.194)
        pf.space_after = Pt(4)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.208


def para(doc, parts, style=None, align=None, space_after=None, space_before=None,
         line=None, keep_next=False):
    """parts: list of (text, bold, italic, color) or plain str."""
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if line is not None:
        p.paragraph_format.line_spacing = line
    if keep_next:
        p.paragraph_format.keep_with_next = True
    if isinstance(parts, str):
        parts = [(parts, False, False, None)]
    for item in parts:
        if isinstance(item, str):
            text, bold, italic, color = item, False, False, None
        else:
            text, bold, italic, color = item
        r = p.add_run(text)
        set_run_font(r, size=11, color=color, bold=bold, italic=italic)
    return p


def bullets(doc, items):
    for item in items:
        if isinstance(item, str):
            parts = [(item, False, False, None)]
        else:
            parts = item
        para(doc, parts, style="List Bullet")


def add_page_number(p):
    run = p.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)
    set_run_font(run, size=9, color=GRAY)


def header_footer_setup(doc):
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run("ShiKe  |  Business Plan")
    set_run_font(r, size=9, color=GRAY)
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D9D9D9")
    pBdr.append(bottom)
    pPr.append(pBdr)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(fp)


def make_table(doc, headers, rows, widths_in, header_fill=HEADER_FILL, font_size=10,
               col_bold=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    tblPr = table._tbl.tblPr

    # Rebuild tblGrid with explicit column widths matching the DXA geometry.
    tbl = table._tbl
    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for w_in in widths_in:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(round(w_in * 1440))))
        grid.append(gc)
    tbl.insert(list(tbl).index(tblPr) + 1, grid)

    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "9360")
    tblW.set(qn("w:type"), "dxa")

    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)

    cellMar = OxmlElement("w:tblCellMar")
    for side, val in (("top", "80"), ("bottom", "80"), ("start", "120"), ("end", "120")):
        el = OxmlElement("w:" + side)
        el.set(qn("w:w"), val)
        el.set(qn("w:type"), "dxa")
        cellMar.append(el)
    tblPr.append(cellMar)

    # header row
    hdr = table.rows[0]
    trPr = hdr._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    trPr.append(tblHeader)
    for j, htext in enumerate(headers):
        cell = hdr.cells[j]
        cell.width = Inches(widths_in[j])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        r = p.add_run(htext)
        set_run_font(r, size=font_size, color=INK, bold=True)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), header_fill)
        cell._tc.get_or_add_tcPr().append(shd)

    for i, row in enumerate(rows):
        tr = table.rows[i + 1]
        for j, val in enumerate(row):
            cell = tr.cells[j]
            cell.width = Inches(widths_in[j])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.line_spacing = 1.05
            if isinstance(val, list):
                parts = val
            else:
                parts = [(str(val), False, False, None)]
            for k, item in enumerate(parts):
                if isinstance(item, str):
                    text, bold, italic, color = item, False, False, None
                else:
                    text, bold, italic, color = item
                if k > 0:
                    r = p.add_run()
                    r.add_break()
                r = p.add_run(text)
                set_run_font(r, size=font_size, color=color, bold=bold, italic=italic)

    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        cant = OxmlElement("w:cantSplit")
        trPr.append(cant)

    return table


def spacing_after_table(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("")
    set_run_font(r, size=4)
    return p


def cover(doc):
    for _ in range(5):
        para(doc, "", space_after=12)
    para(doc, [("SHIKE", False, False, INK)], align=WD_ALIGN_PARAGRAPH.CENTER,
         space_after=4, line=1.0)
    for r in doc.paragraphs[-1].runs:
        set_run_font(r, size=40, color=INK, bold=True)
    para(doc, [("The AI food scanner that turns any refrigerator into a smart pantry",
                False, True, DARK)], align=WD_ALIGN_PARAGRAPH.CENTER, space_after=34, line=1.1)
    for r in doc.paragraphs[-1].runs:
        set_run_font(r, size=16, color=DARK, italic=True)
    para(doc, [("BUSINESS PLAN", False, False, INK)], align=WD_ALIGN_PARAGRAPH.CENTER,
         space_after=6, line=1.0)
    for r in doc.paragraphs[-1].runs:
        set_run_font(r, size=20, color=INK, bold=True)
    para(doc, [("Prepared for SSG Accelerator  ·  Cohort 3 2026", False, False, GRAY)],
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    for r in doc.paragraphs[-1].runs:
        set_run_font(r, size=11, color=GRAY)
    para(doc, [("August 2026  ·  Confidential  ·  Version 1.0", False, False, GRAY)],
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)
    for r in doc.paragraphs[-1].runs:
        set_run_font(r, size=11, color=GRAY)
    para(doc, [("Contact: ", False, True, GRAY), ("<Founder name>  ·  <Email>  ·  <Phone>",
                False, False, GRAY)], align=WD_ALIGN_PARAGRAPH.CENTER)
    for r in doc.paragraphs[-1].runs:
        set_run_font(r, size=10, color=GRAY)
    doc.add_page_break()


def main():
    doc = Document()
    style_setup(doc)
    header_footer_setup(doc)
    cover(doc)

    # 1 Executive summary
    doc.add_heading("1  Executive Summary", level=1)
    para(doc, "ShiKe is an AI food scanner that clips onto any existing refrigerator and turns "
              "it into a smart pantry. Users pass food in front of the scanner once; on-device "
              "recognition (barcode-first, with vision fallback) identifies the item and its "
              "expiry date, and the family inventory updates automatically — no typing, no "
              "manual barcode hunting, no tags. "
              "A fridge-side screen, a mobile app, and a home dashboard then answer the three "
              "questions every household asks: What do we have? What is about to expire? What "
              "can we cook tonight?")
    para(doc, "One-third of the world's food is wasted. China discards more food than any other "
              "country (about 108 million tons in 2024), and Japan loses 4.6 million tons a year, "
              "2.24 million of them in households. Yet the roughly 324 million urban fridges in "
              "China have no affordable way to manage what is inside: smart fridges only reach "
              "new buyers, and manual-entry apps fail because people forget to type.")
    para(doc, "ShiKe's answer is a low-cost retrofit with zero learning curve. We have already "
              "shipped the full loop — hardware to cloud AI to software — as a working prototype: "
              "a dual-module hardware prototype, a live backend (shike.live, 11 APIs, 15 tables), mobile H5, "
              "web dashboard, and a 3D-printed enclosure. We are raising an 18-month seed round "
              "to push recognition accuracy past 90%, launch through crowdfunding, and enter "
              "Japan — the first market with policy tailwind, standardized date labels, and no "
              "competing scanner product.")

    snap = make_table(
        doc,
        ["", "Snapshot"],
        [
            ["Product", "AI food scanner + family inventory platform for existing refrigerators"],
            ["Status", "Working end-to-end prototype; live backend; V1 production-ready design in build"],
            ["Market", "China: TAM ¥20–40B, SAM ¥2–4B; Japan first global market (food-loss policy leader)"],
            ["Business model", "Hardware (¥299) + subscription (¥99/yr) + B2B expiry-compliance SaaS"],
            ["Team", "Four founders covering hardware, cloud AI, product, and go-to-market"],
            ["Raise", "¥2.5–4.0M (≈US$350–560K) seed for an 18-month runway"],
        ],
        [1.5, 5.0],
    )
    spacing_after_table(doc)

    # 2 Problem
    doc.add_heading("2  The Problem: Food Invisible in the Fridge", level=1)
    para(doc, "Food waste is one of the largest avoidable costs in the modern household, and its "
              "root cause is visibility, not cooking skill. Once food goes into a refrigerator, "
              "it leaves the family's attention until it is too late.")
    bullets(doc, [
        [("Global scale: ", True, False, None),
         ("roughly one-third of all food produced is wasted; China was the world's largest "
          "food-wasting country in 2024 at about 108 million tons per year (≈76 kg per person), "
          "and the restaurant sector alone wastes 17–18 million tons a year.", False, False, None)],
        [("Japan's policy target: ", True, False, None),
         ("FY2024 food loss was 4.61 million tons, 2.24 million from households; unopened "
          "“direct disposal” is a named target of the Food Loss Reduction Act. Single-person "
          "households are 40% of all homes, and seniors waste the most (46 kg/person/year in "
          "their 70s vs. 16.6 kg under 30).", False, False, None)],
        [("Why it happens: ", True, False, None),
         ("the fridge is out of sight, out of mind — expiry dates are discovered too late, "
          "food is rebought, and spoiled items are thrown away unopened.", False, False, None)],
        [("Existing answers fail: ", True, False, None),
         ("built-in smart fridges cost thousands of RMB and only reach new buyers; manual-entry "
          "apps demand ~30 seconds per item and die from friction; per-item tags add cost and "
          "effort to every single product.", False, False, None)],
    ])

    # 3 Solution & product
    doc.add_heading("3  The Solution: Scan, Recognize, Record, Display", level=1)
    para(doc, "ShiKe compresses food management into a one-second action. The core loop is "
              "“scan → recognize → record → display”: a device mounted on the fridge photographs "
              "each item, cloud vision AI identifies the name and expiry date, the inventory "
              "updates by ±1, and every screen in the home reflects the result immediately.")
    bullets(doc, [
        [("Hardware scanner: ", True, False, None),
         ("dual-module device — an outer touchscreen tablet plus an inner camera module — "
          "works on any fridge, cabinet, or medicine drawer.", False, False, None)],
        [("Cloud AI recognition: ", True, False, None),
         ("identifies product name and expiry date in milliseconds to under two seconds; "
          "low-confidence results are rejected and the user is asked to rescan rather than "
          "guessing.", False, False, None)],
        [("Family inventory: ", True, False, None),
         ("stock tracked by container (fridge, snack box, medicine box, spice rack, pantry), "
          "with restock and consumption modes.", False, False, None)],
        [("Proactive alerts: ", True, False, None),
         ("expiring-soon reminders 3 days ahead, on the device, the app, and the dashboard.", False, False, None)],
        [("Recipes & shopping: ", True, False, None),
         ("“what can I cook with what I have” recommendations and an auto-generated shopping "
          "list for missing items.", False, False, None)],
        [("Family sharing & waste stats: ", True, False, None),
         ("multi-member sync, remote visibility for children checking parents' fridges, and "
          "monthly waste/savings reports.", False, False, None)],
    ])
    para(doc, "Three everyday scenarios define the product: batch restocking (scan every item "
              "while putting groceries in), single-item in/out (grab a snack, scan it once), and "
              "standby display (a second screen on the fridge showing what expires next and what "
              "to cook tonight).")

    # 4 Technology & PoC
    doc.add_heading("4  Technology and Proof of Concept", level=1)
    para(doc, "ShiKe is not a slideware concept. The complete loop is live today and has been "
              "demonstrated end to end:")
    bullets(doc, [
        [("Hardware: ", True, False, None),
         ("a dual-module build — an outer T5-E1 touchscreen tablet and an inner XIAO ESP32-S3 "
          "camera module with barcode-first scan and voice feedback — running TuyaOpen + "
          "Arduino firmware with live HTTPS upload.", False, False, None)],
        [("Cloud: ", True, False, None),
         ("a live backend at shike.live (Next.js + Prisma + SQLite on Docker) with 11 APIs and "
          "15 tables, including real scan data.", False, False, None)],
        [("Vision AI: ", True, False, None),
         ("local-first pipeline — on-device EAN-13 barcode decode (ms), on-device YOLO "
          "category (≤500 ms), cloud lookup for exact name/expiry via Doubao (P95 <2 s); "
          "low-confidence results not recorded.", False, False, None)],
        [("Software: ", True, False, None),
         ("mobile H5 (5 tabs), a web dashboard, and an interactive hardware demo; the H5 scan "
          "page calls the real API.", False, False, None)],
        [("Design: ", True, False, None),
         ("3D-printed house-shaped enclosure (white body, red roof, magnetic back cover).", False, False, None)],
    ])
    para(doc, "The V1 production-oriented design is already specified and in build: an inner "
              "module (XIAO ESP32-S3 Sense with camera, microphone, light-sensor door detection, "
              "LED, and scan button) works with an outer touchscreen tablet to enable door-side "
              "continuous scanning, barcode decoding, voice feedback, and an offline queue that "
              "syncs when Wi-Fi returns. Device-token authentication and migration to AWS "
              "(S3, DynamoDB/RDS, Bedrock/Rekognition, API Gateway + Lambda, Cognito, CloudFront, "
              "Tokyo region) are planned for production and Japan data compliance (APPI).")
    para(doc, [("Honest gap and plan: ", True, False, None),
               ("recognition and date-extraction accuracy is the make-or-break metric. We are "
                "running a 100-user pilot, building a 1,000+ image test set, adding date-OCR "
                "post-processing, enforcing “don't guess” rules, and closing the loop with user "
                "corrections feeding back into the knowledge base.", False, False, None)])

    # 5 Market
    doc.add_heading("5  Market Opportunity", level=1)
    para(doc, "The market is anchored in a real, measured social problem with strong policy "
              "tailwind in our first markets:")
    make_table(
        doc,
        ["Level", "Definition", "Size"],
        [
            ["TAM", "China household food-management market",
             "¥20–40B/year (324M urban fridges × ¥60–120 yearly spend)"],
            ["SAM", "Retrofit smart accessory + subscription segment",
             "¥2–4B/year (smart-fridge market ¥2.16–2.35B in 2025; new external-accessory demand)"],
            ["SOM", "ShiKe's reachable share in 3 years",
             "¥30–100M/year (50–200K units × ¥250–350 incl. subscription)"],
        ],
        [1.1, 2.8, 2.6],
    )
    spacing_after_table(doc)
    bullets(doc, [
        [("China: ", True, False, None),
         ("104 urban fridges per 100 households (~324M units); 30–39M replaced per year; fridge "
          "retail ~¥113B in 2025 — the installed base dwarfs annual new purchases.", False, False, None)],
        [("Global: ", True, False, None),
         ("smart kitchen appliances ≈US$21B (2023), projected ≈US$43B by 2030 (CAGR 10.5%).", False, False, None)],
        [("Tailwinds: ", True, False, None),
         ("LLM vision cost/accuracy improving fast; pre-made food and fresh e-commerce "
          "proliferation increases item variety; the silver-economy “children watching parents' "
          "fridge” scenario is rising.", False, False, None)],
    ])

    # 6 Competition
    doc.add_heading("6  Competition and Differentiation", level=1)
    para(doc, "ShiKe's wedge is the only position that covers the fridges already in people's "
              "homes: “spend a few hundred RMB and your old fridge becomes smart.”")
    make_table(
        doc,
        ["Category", "Representative players", "Why ShiKe wins"],
        [
            ["Built-in smart fridges",
             "Haier AI Eye, Midea, Samsung Bespoke",
             "Require replacing a fridge (¥10K+); reach only ~7–8% of units sold per year; date "
             "reading remains the weak spot. ShiKe covers 100% of installed fridges."],
            ["Manual-entry apps",
             "NoWaste, CozZo, 摆烂冰箱",
             "Free but ~30 s/item with high drop-off. ShiKe removes typing entirely — one scan "
             "with a physical touchpoint next to the fridge."],
            ["Tag / clip solutions",
             "Ovie LightTags, Midea NFC clips",
             "Per-item tags add cost and effort. ShiKe changes nothing about the food itself and "
             "supports batch scanning."],
            ["Canteen / restaurant SaaS",
             "Back-office inventory suites",
             "Process-heavy, no vision. Camera-based expiry ledger with automatic reports is a "
             "differentiation module."],
            ["Platform apps",
             "Hema, Meituan grocery",
             "Platform-locked and non-neutral. ShiKe is hardware-anchored and independent."],
        ],
        [1.4, 2.0, 3.1],
    )
    spacing_after_table(doc)
    para(doc, "We also study the cautionary tale of Ovie Smarterware (US food-waste hardware): "
              "moral appeal alone does not pay, hardware-plus-subscription cash cycles are slow, "
              "and a single form factor can be commoditized. ShiKe therefore anchors value in "
              "money and time saved (“how much you stopped throwing away”), pairs hardware with "
              "subscription revenue, and keeps a B2B and licensing path open.")

    # 7 Business model
    doc.add_heading("7  Business Model and Unit Economics", level=1)
    para(doc, "Revenue is built on three lines: consumer hardware, a consumer subscription, and "
              "a B2B compliance SaaS — all sharing one recognition and inventory core.")
    make_table(
        doc,
        ["Revenue stream", "3-yr target share", "Description"],
        [
            ["C-end hardware", "≈45%",
             "Scanner at ¥299; family bundle (scanner + medicine-box edition) at ¥499"],
            ["C-end subscription", "≈25%",
             "ShiKe Cloud Pro: expiry push, recipes, family sharing, nutrition reports — ¥9.9/mo "
             "or ¥99/yr, first year free"],
            ["B2B SaaS", "≈25%",
             "Expiry ledger, waste reports, batch inventory — ¥299–999/store/year"],
            ["Other (API / data partnerships)", "≈5%",
             "Recognition API licensing to appliance and SaaS vendors"],
        ],
        [1.7, 1.2, 3.6],
    )
    spacing_after_table(doc)
    para(doc, "Unit economics per user over three years:")
    make_table(
        doc,
        ["Metric", "Value", "Note"],
        [
            ["Hardware revenue", "¥299", "one-time"],
            ["Subscription revenue (3 yr)", "≈¥129", "first year free; retention 60%/40%/30%"],
            ["Total revenue per user", "≈¥428", "3-year view"],
            ["Hardware + cloud + fulfillment cost", "¥130–160", "incl. 3 years of vision calls"],
            ["Marketing CAC target", "≤¥90", "content-driven, not paid-scale"],
            ["3-year gross profit", "≈¥180–200", "LTV/CAC ≥ 3"],
        ],
        [2.0, 1.2, 3.3],
    )
    spacing_after_table(doc)
    para(doc, "BOM cost drops from ¥110–175 (prototype) to ¥72–110 (mass production); hardware "
              "gross margin is ≈70% at ¥299 vs. ¥90 COGS, and the blended gross-margin target is "
              "≥55% after cloud-vision costs. Crowdfunding also pre-funds inventory: 5,000 units "
              "at ¥299 ≈ ¥1.5M of working capital.")

    # 8 GTM & expansion
    doc.add_heading("8  Go-to-Market and Global Expansion", level=1)
    para(doc, "Cold start (0–6 months):", )
    bullets(doc, [
        [("100-user seed program", True, False, None),
         (" — free units for real usage data; gates: recognition ≥90%, ≥8 active scans/month, "
          "recommendation score ≥8/10.", False, False, None)],
        [("Content seeding", True, False, None),
         (" — Xiaohongshu/Douyin emotional topics (“3-year-old can found in the fridge”, "
          "“grandma's fridge”), KOC distribution.", False, False, None)],
        [("Crowdfunding", True, False, None),
         (" — 2,000–5,000 units on Xiaomi Youpin/JD crowdfunding, validating C-end willingness "
          "to pay and returning cash.", False, False, None)],
        [("B2B pilots", True, False, None),
         (" — 10 canteen/restaurant stores free trial, ≥30% conversion to paid.", False, False, None)],
    ])
    para(doc, "Scaling (6–24 months): Tmall/JD flagship stores plus KOL distribution; "
              "city-by-city B2B expansion; appliance-retailer “upgrade your old fridge” corners; "
              "and overseas entry.")
    para(doc, [("Global expansion — Japan first. ", True, False, None),
               ("Japan has the strongest policy and social consensus on food loss, a precise pain "
                "segment (40% single-person households), standardized 消費期限/賞味期限 labels "
                "that suit AI OCR, and no competing scanner-form-factor product. Route: Makuake "
                "crowdfunding for validation → Amazon Japan + Rakuten, with PSE/TELEC/METI "
                "certification and AWS Tokyo deployment for APPI compliance. The US follows as "
                "the second market with the Japan-proven version; the EU is reassessed after the "
                "2027 CRA requirements are clarified.", False, False, None)])

    # 9 Financial plan
    doc.add_heading("9  Financial Plan", level=1)
    para(doc, "Assumptions: C-end average ¥299 per unit; B-end ¥199/unit + ¥499/year SaaS; "
              "cloud and fulfillment costs in COGS; four-person team with annual operating cost "
              "of roughly ¥2–4M. Three scenarios:")
    make_table(
        doc,
        ["Scenario", "Year 1", "Year 2", "Year 3", "3-yr revenue", "3-yr gross profit"],
        [
            ["Conservative", "3K units + 60 B-end", "12K + 200", "30K + 500", "≈¥20M", "≈¥10M"],
            ["Base", "8K units + 150 B-end", "40K + 600", "120K + 1,500", "≈¥70M", "≈¥38M"],
            ["Optimistic", "20K units + 400 B-end", "100K + 1,500", "300K + 4,000", "≈¥180M", "≈¥98M"],
        ],
        [0.9, 1.3, 1.0, 1.3, 1.0, 1.0],
    )
    spacing_after_table(doc)
    bullets(doc, [
        [("Year 1 is achievable without miracles: ", True, False, None),
         ("8,000 units ≈ one successful crowdfunding campaign plus 150 B-end stores in one city.", False, False, None)],
        [("The 8K → 40K jump depends on ", True, False, None),
         ("recognition accuracy, subscription renewal ≥40%, and 1–2 channels (e-commerce, "
          "appliance retail, or chain catering).", False, False, None)],
        [("Base case of ≈¥70M in 3 years ", True, False, None),
         ("supports a 5–8× seed-valuation narrative for an early hardware team.", False, False, None)],
    ])

    # 10 Team
    doc.add_heading("10  Team", level=1)
    para(doc, "ShiKe is built by a four-person founding team that covers the full skill set for "
              "hardware-startup execution:")
    bullets(doc, [
        [("Embedded & hardware: ", True, False, None),
         ("TuyaOpen firmware on T5-E1, ESP32-S3 camera modules, wiring, flashing, and device QA.", False, False, None)],
        [("Cloud AI & backend: ", True, False, None),
         ("Next.js, Prisma, Docker; 11 live APIs and a 15-table schema; vision pipeline integration.", False, False, None)],
        [("Frontend & product: ", True, False, None),
         ("mobile H5, web dashboard, interactive hardware demo, and interaction design.", False, False, None)],
        [("Product narrative & GTM: ", True, False, None),
         ("pitch and story, demo direction, market research (China + Japan), and commercialization "
          "planning.", False, False, None)],
    ])
    para(doc, "Execution evidence: the team converged the product definition on one board, then "
              "shipped the end-to-end prototype in roughly 48 hours (Aug 14 evening → Aug 16 noon) "
              "with disciplined SOPs — API contracts, flashing procedures, test plans, BOM, and "
              "acceptance gates. This combination of speed and documentation is the team's "
              "strongest asset for an accelerator program.")

    # 11 Risks
    doc.add_heading("11  Risks and Mitigation", level=1)
    make_table(
        doc,
        ["Risk", "Severity", "Mitigation"],
        [
            ["Recognition accuracy below 90%",
             "High",
             "30-day focus: multi-angle dataset, date-OCR post-processing, “don't guess” rules, "
             "user-correction feedback loop; 100-user pilot as gate"],
            ["Giants bundle the feature free",
             "High",
             "No head-on war: focus on installed fridges, non-fridge containers (medicine boxes, "
             "snack boxes), and the children-parents narrative"],
            ["Weak willingness to pay",
             "Medium-high",
             "Sell money/time saved (“how much you stopped wasting”), not morality; subscription "
             "anchored on push/recipes/sharing"],
            ["Hardware supply chain & certification",
             "Medium",
             "Crowdfunding-first, small-batch manufacturing, certification planned in parallel"],
            ["Privacy & compliance (PIPL/APPI)",
             "Medium",
             "Privacy-first statements, local processing options, camera kill-switch, Tokyo region "
             "for Japan data"],
            ["Cash-flow pressure",
             "Medium",
             "Subscription upfront, rolling production batches, B2B prepayments"],
            ["Single-vendor dependency",
             "Medium-low",
             "Model abstraction layer; domain/server migrated to company ownership before scale"],
        ],
        [2.0, 1.1, 3.4],
    )
    spacing_after_table(doc)

    # 12 Roadmap
    doc.add_heading("12  Roadmap (0–24 Months)", level=1)
    make_table(
        doc,
        ["Phase", "Time", "Goal", "Key milestones"],
        [
            ["Validation", "0–3 mo",
             "Prove product-market fit",
             "Recognition ≥90%; 100 seed users; retention data; company registration"],
            ["Crowdfunding", "3–6 mo",
             "Validate C-end payment & cash flow",
             "2,000–5,000 units funded; mass-production BOM and certification started; 10 B-end pilots"],
            ["Scaling", "6–12 mo",
             "Build channels & renewal",
             "E-commerce launch; 150 B-end stores; subscription renewal ≥40%; gen-2 medicine/snack "
             "box edition"],
            ["Expansion", "12–24 mo",
             "Scale & second curve",
             "Chain B2B; Japan evaluation/launch; recognition-API licensing; break-even at 100K "
             "cumulative units"],
        ],
        [1.0, 0.8, 1.7, 3.0],
    )
    spacing_after_table(doc)

    # 13 Ask
    doc.add_heading("13  The Ask and Use of Funds", level=1)
    para(doc, "We are raising ¥2.5–4.0M (≈US$350–560K) for an 18-month runway, either as a "
              "standalone seed or combined with accelerator support, crowdfunding proceeds "
              "(≈¥1.5M from 5,000 units), and founder capital. Beyond capital, we are seeking "
              "go-to-market mentorship for Japan entry, supply-chain and certification guidance, "
              "and access to a global mentor and investor network — exactly what an AI-focused "
              "global accelerator provides.")
    make_table(
        doc,
        ["Use of funds", "Share", "Amount", "Purpose"],
        [
            ["Product & recognition", "25%", "¥0.6–1.0M", "Vision dataset, model tuning, firmware iteration"],
            ["Manufacturing & certification", "30%", "¥0.8–1.2M", "Molds, first 5,000 units, CCC/SRRC certification"],
            ["Go-to-market", "25%", "¥0.6–1.0M", "Crowdfunding, KOC, B2B sales"],
            ["Team & operations", "20%", "¥0.5–0.8M", "4-person team, 18 months"],
        ],
        [1.9, 0.8, 1.1, 2.7],
    )
    spacing_after_table(doc)
    para(doc, [("Note: ", True, False, None),
               ("all figures are in RMB (¥); approximate conversion at 7.2 RMB/USD. Contact "
                "details and founder names are intentionally left as placeholders pending team "
                "confirmation.", False, False, GRAY)])
    for r in doc.paragraphs[-1].runs:
        set_run_font(r, size=10, color=GRAY)

    out = "/Users/lwl/Documents/ChatGPT/HsHH/docs/04-商业/ShiKe-Business-Plan.docx"
    doc.save(out)
    print("saved", out)


if __name__ == "__main__":
    main()
